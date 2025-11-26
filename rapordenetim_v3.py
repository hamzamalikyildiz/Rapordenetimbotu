import os
import json
import sqlite3
from datetime import datetime
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ApplicationBuilder, ContextTypes, CommandHandler, MessageHandler, filters, ConversationHandler, CallbackQueryHandler
from docx import Document
import google.generativeai as genai
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

if not os.path.exists('credentials.json'):
    if 'CREDENTIALS_JSON' in os.environ:
        with open('credentials.json', 'w') as f:
            f.write(os.environ['CREDENTIALS_JSON'])

if not os.path.exists('token.json'):
    if 'TOKEN_JSON' in os.environ:
        with open('token.json', 'w') as f:
            f.write(os.environ['TOKEN_JSON'])

TELEGRAM_TOKEN = os.environ.get('TELEGRAM_TOKEN', "8034942082:AAER34bDTUszcntdE6ncavK1xohRJ4bNIcE")
GOOGLE_API_KEY = os.environ.get('GOOGLE_API_KEY', "AIzaSyDWTC1ScC7mfkz3tD3Qmyl6LFNdVIo9VnM")
DRIVE_KLASOR_ID = "10X9Qd9o1vTR2IJ3ZuhjPBS2U4JacKjr6"
SCOPES = ['https://www.googleapis.com/auth/drive']

genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-2.0-flash')

TARIH_ADIMI, PROJELER_ADIMI, TAMAMLANAN_GOREVLER_ADIMI, DEVAM_EDEN_GOREVLER_ADIMI, SORUNLAR_ADIMI, GELECEK_PLANLAR_ADIMI, EK_NOTLAR_ADIMI = range(7)

def veritabani_kur():
    conn = sqlite3.connect('raporlar.db')
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS detayli_raporlar (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            tarih_araligi TEXT,
            projeler TEXT,
            tamamlanan_gorevler TEXT,
            devam_eden_gorevler TEXT,
            sorunlar TEXT,
            gelecek_planlar TEXT,
            ek_notlar TEXT
        )
    ''')
    conn.commit()
    conn.close()

def drive_yukle(dosya_yolu):
    creds = None
    if os.path.exists('token.json'):
        try:
            creds = Credentials.from_authorized_user_file('token.json', SCOPES)
        except Exception:
            pass
    
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
                with open('token.json', 'w') as token:
                    token.write(creds.to_json())
            except Exception:
                return None
        else:
            if not os.path.exists('credentials.json'):
                return None
            try:
                flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
                creds = flow.run_local_server(port=0)
                with open('token.json', 'w') as token:
                    token.write(creds.to_json())
            except:
                return None

    try:
        service = build('drive', 'v3', credentials=creds)
        file_metadata = {'name': os.path.basename(dosya_yolu)}
        if DRIVE_KLASOR_ID:
            file_metadata['parents'] = [DRIVE_KLASOR_ID]
            
        media = MediaFileUpload(dosya_yolu, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        file = service.files().create(body=file_metadata, media_body=media, fields='id, webViewLink').execute()
        
        file_id = file.get('id')
        permission = {
            'type': 'anyone',
            'role': 'reader',
        }
        service.permissions().create(
            fileId=file_id,
            body=permission,
            fields='id',
        ).execute()

        return file.get('webViewLink')
    except Exception:
        return None
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    butonlar = [[
        InlineKeyboardButton("Yeni Rapor Ekle", callback_data='yeni_rapor'),
        InlineKeyboardButton("Iptal et", callback_data="iptal_et")
    ]]
    panel = InlineKeyboardMarkup(butonlar)
    await context.bot.send_message(
        chat_id=update.effective_chat.id, 
        text="Merhaba. Haftalik rapor olusturucuya hos geldin. Ne yapmak istersiniz?\nYardim icin: /yardim", 
        reply_markup=panel
    )

async def yardim(update: Update, context: ContextTypes.DEFAULT_TYPE):
    mesaj = """
RAPOR BOTU KULLANIM KILAVUZU

Bu bot, haftalik raporlarini hizlica olusturup Google Drive'a yuklemeni saglar.

KOMUTLAR:
/start - Ana menuyu acar ve botu baslatir.
/yardim - Bu menuyu gosterir.
/iptal - Rapor olusturma surecini iptal eder.

NASIL KULLANILIR?
1. /start yazarak menuyu ac.
2. 'Yeni Rapor Ekle' butonuna tikla.
3. Tarih araligini gir (Orn: 20-27 Kasim).
4. Botun sorularina cevap ver.
5. Bot, cevaplarini yapay zeka ile duzenleyip rapor haline getirir ve Drive linkini sana verir.
    """
    await context.bot.send_message(chat_id=update.effective_chat.id, text=mesaj)

async def butona_tiklandi(update: Update, context: ContextTypes.DEFAULT_TYPE):
    tiklama = update.callback_query
    await tiklama.answer()
    
    if tiklama.data == 'yeni_rapor':
        await tiklama.edit_message_text(text="Önce tarihi ayarlayalım. Rapor hangi tarihleri kapsıyor? (Orn: 20-27 Kasım 2025)")
        return TARIH_ADIMI
    elif tiklama.data == 'iptal_et':
        await tiklama.edit_message_text(text="İşlem iptal edildi.")
        return ConversationHandler.END

async def tarih_al(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['tarih_araligi'] = update.message.text
    await update.message.reply_text(f"Tarih: {update.message.text}\n\n1. Bu hafta hangi PROJELER uzerinde calistin?")
    return PROJELER_ADIMI

async def projeler_al(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['projeler'] = update.message.text
    await update.message.reply_text("2. TAMAMLANAN görevlerin neler?")
    return TAMAMLANAN_GOREVLER_ADIMI

async def tamamlanan_gorevler_al(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['tamamlanan_gorevler'] = update.message.text
    await update.message.reply_text("3. Şu an DEVAM EDEN işler neler?")
    return DEVAM_EDEN_GOREVLER_ADIMI

async def devam_eden_gorevler_al(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['devam_eden_gorevler'] = update.message.text
    await update.message.reply_text("4. Karşılaştığın SORUNLAR veya engeller var mı?")
    return SORUNLAR_ADIMI

async def sorunlar_al(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['sorunlar'] = update.message.text
    await update.message.reply_text("5. GELECEK HAFTA için planların neler?")
    return GELECEK_PLANLAR_ADIMI

async def gelecek_planlar_al(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['gelecek_planlar'] = update.message.text
    await update.message.reply_text("6. Eklemek istediğin NOTLAR var mi? (Yoksa 'Yok' yazabilirsin)")
    return EK_NOTLAR_ADIMI

async def ek_notlar_al(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['ek_notlar'] = update.message.text
    await update.message.reply_text("Veriler alindi. AI her basligi tek tek düzenliyor, rapor hazırlanıyor...")

    veri = context.user_data
    kullanici_adi = update.effective_user.first_name.replace(" ", "_")
    bugun = datetime.now().strftime("%d-%m-%Y")
    dinamik_dosya_adi = f"Rapor_{kullanici_adi}_{bugun}.docx"

    try:
        conn = sqlite3.connect('raporlar.db')
        c = conn.cursor()
        c.execute('''
            INSERT INTO detayli_raporlar (user_id, tarih_araligi, projeler, tamamlanan_gorevler, devam_eden_gorevler, sorunlar, gelecek_planlar, ek_notlar)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            update.effective_user.id, veri['tarih_araligi'], veri['projeler'], veri['tamamlanan_gorevler'], 
            veri['devam_eden_gorevler'], veri['sorunlar'], veri['gelecek_planlar'], veri['ek_notlar']
        ))
        conn.commit()
        conn.close()
    except Exception:
        pass

    duzenlenmis_veri = {}
    try:
        prompt = f"""
        Asagidaki ham rapor verilerini analiz et. Her bir baslik altindaki metni profesyonel, kurumsal bir dille yeniden yaz.
        
        HAM VERILER:
        - Projeler: {veri['projeler']}
        - Tamamlanan: {veri['tamamlanan_gorevler']}
        - Devam Eden: {veri['devam_eden_gorevler']}
        - Sorunlar: {veri['sorunlar']}
        - Planlar: {veri['gelecek_planlar']}
        - Notlar: {veri['ek_notlar']}
        
        CIKTI FORMATI (Sadece bu JSON'u ver):
        {{
            "projeler": "Duzenlenmis metin...",
            "tamamlanan": "Duzenlenmis metin...",
            "devam_eden": "Duzenlenmis metin...",
            "sorunlar": "Duzenlenmis metin...",
            "planlar": "Duzenlenmis metin...",
            "notlar": "Duzenlenmis metin...",
            "ozet": "Kisa bir yonetici ozeti..."
        }}
        """
        response = model.generate_content(prompt)
        json_str = response.text.strip()
        if json_str.startswith("```json"):
            json_str = json_str[7:-3]
        duzenlenmis_veri = json.loads(json_str)
    except Exception:
        duzenlenmis_veri = {
            "projeler": veri['projeler'],
            "tamamlanan": veri['tamamlanan_gorevler'],
            "devam_eden": veri['devam_eden_gorevler'],
            "sorunlar": veri['sorunlar'],
            "planlar": veri['gelecek_planlar'],
            "notlar": veri['ek_notlar'],
            "ozet": "AI Duzenlemesi yapilamadi."
        }

    belge = Document()
    belge.add_heading('HAFTALIK FAALIYET RAPORU', 0)
    p = belge.add_paragraph()
    p.add_run(f"Rapor Sahibi: {update.effective_user.first_name}\n").bold = True
    p.add_run(f"Tarih Araligi: {veri['tarih_araligi']}").bold = True
    
    belge.add_heading('Yonetici Ozeti', level=1)
    belge.add_paragraph(duzenlenmis_veri.get('ozet', ''))
    
    basliklar_sirasi = [
        ('Projeler', duzenlenmis_veri.get('projeler', '')),
        ('Tamamlanan Gorevler', duzenlenmis_veri.get('tamamlanan', '')),
        ('Devam Eden Isler', duzenlenmis_veri.get('devam_eden', '')),
        ('Sorunlar ve Engeller', duzenlenmis_veri.get('sorunlar', '')),
        ('Gelecek Hafta Plani', duzenlenmis_veri.get('planlar', '')),
        ('Ek Notlar', duzenlenmis_veri.get('notlar', ''))
    ]
    
    for baslik, icerik in basliklar_sirasi:
        belge.add_heading(baslik, level=1)
        p = belge.add_paragraph(style='List Bullet')
        p.add_run(icerik)

    belge.save(dinamik_dosya_adi)
    
    drive_linki = drive_yukle(dinamik_dosya_adi)
    
    mesaj = f"Rapor hazir: {dinamik_dosya_adi}\n"
    if drive_linki:
        mesaj += f"Drive Linki: {drive_linki}"
    else:
        mesaj += "Drive'a yuklenirken hata olustu."

    await context.bot.send_document(chat_id=update.effective_chat.id, document=open(dinamik_dosya_adi, 'rb'))
    await context.bot.send_message(chat_id=update.effective_chat.id, text=mesaj)
    
    if os.path.exists(dinamik_dosya_adi):
        os.remove(dinamik_dosya_adi)
        
    return ConversationHandler.END

async def iptal(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Islem iptal edildi.")
    return ConversationHandler.END

if __name__ == '__main__':
    veritabani_kur()
    
    uygulama = ApplicationBuilder().token(TELEGRAM_TOKEN).build()
    
    sohbet_yoneticisi = ConversationHandler(
        entry_points=[
            CallbackQueryHandler(butona_tiklandi, pattern='^yeni_rapor$'),
            CallbackQueryHandler(butona_tiklandi, pattern='^iptal_et$')
        ],
        states={
            TARIH_ADIMI: [MessageHandler(filters.TEXT & ~filters.COMMAND, tarih_al)],
            PROJELER_ADIMI: [MessageHandler(filters.TEXT & ~filters.COMMAND, projeler_al)],
            TAMAMLANAN_GOREVLER_ADIMI: [MessageHandler(filters.TEXT & ~filters.COMMAND, tamamlanan_gorevler_al)],
            DEVAM_EDEN_GOREVLER_ADIMI: [MessageHandler(filters.TEXT & ~filters.COMMAND, devam_eden_gorevler_al)],
            SORUNLAR_ADIMI: [MessageHandler(filters.TEXT & ~filters.COMMAND, sorunlar_al)],
            GELECEK_PLANLAR_ADIMI: [MessageHandler(filters.TEXT & ~filters.COMMAND, gelecek_planlar_al)],
            EK_NOTLAR_ADIMI: [MessageHandler(filters.TEXT & ~filters.COMMAND, ek_notlar_al)],
        },
        fallbacks=[CommandHandler('start', start)]
    )
    
    uygulama.add_handler(CommandHandler('start', start))
    uygulama.add_handler(CommandHandler('yardim', yardim))
    uygulama.add_handler(sohbet_yoneticisi)
    
    print("Bot calisiyor...")
    uygulama.run_polling()



